"""DOCX extractor using python-docx. Paragraph + easily-available table text."""
from __future__ import annotations

from typing import List

from models import TextUnit


def extract(file_path: str) -> List[TextUnit]:
    import docx  # python-docx; imported lazily

    document = docx.Document(file_path)
    parts: List[str] = []

    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)

    # Pull plain text out of tables when it's readily available; ignore images.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    return [TextUnit(text=text, page_number=None)]
